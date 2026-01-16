from docx import Document

def export_docx(notes: str, output_path="lecture_notes.docx"):
    doc = Document()
    doc.add_heading("Lecture Notes", level=1)

    for line in notes.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)
    return output_path
